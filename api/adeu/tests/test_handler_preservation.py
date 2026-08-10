"""
Preservation tests for the deprecated api/adeu Vercel handler.

The handler is retained per ADR-004 §4 (its authors have not approved
removal). These tests pin its observable behavior — all five routes, error
shapes, the 50 MB Content-Length pre-check — and verify that its shared
schemas are now imported from services/document-editor, the authoritative
implementation, rather than from the removed sidecar/.
"""

import base64
import io
import json

from docx import Document

from .conftest import multipart_request


def _json_body(response: dict) -> dict:
    return json.loads(response["body"])


def _docx_body(response: dict) -> Document:
    assert response.get("isBase64Encoded") is True
    return Document(io.BytesIO(base64.b64decode(response["body"])))


# ── Schema wiring: services/document-editor is the source of truth ──────────

class TestSchemaImportWiring:
    def test_schemas_come_from_document_editor_service(self, adeu_index):
        module_file = adeu_index.ProcessBatchRequest.__module__
        assert module_file == "app.schemas.adeu"
        import sys

        loaded_from = sys.modules["app.schemas.adeu"].__file__
        assert "services" in loaded_from and "document-editor" in loaded_from, (
            f"schemas must load from services/document-editor, got {loaded_from}"
        )

    def test_sidecar_is_not_referenced_in_code(self, adeu_index):
        """No import or path may point at the removed sidecar/ directory
        (prose mentions of the history in comments are fine)."""
        import inspect

        source = inspect.getsource(adeu_index)
        assert "from sidecar" not in source
        assert '"sidecar"' not in source and "'sidecar'" not in source


# ── /api/adeu/read ──────────────────────────────────────────────────────────

class TestRead:
    def test_read_valid_docx_returns_text(self, handler, simple_docx):
        resp = handler(
            multipart_request(
                "/api/adeu/read",
                {"file": ("test.docx", simple_docx), "clean_view": "false"},
            )
        )
        assert resp["statusCode"] == 200
        body = _json_body(resp)
        assert "Hello world" in body["text"]
        assert body["filename"] == "test.docx"

    def test_read_invalid_docx_returns_422(self, handler):
        resp = handler(
            multipart_request(
                "/api/adeu/read", {"file": ("bad.docx", b"not a docx file")}
            )
        )
        assert resp["statusCode"] == 422
        assert "detail" in _json_body(resp)


# ── /api/adeu/process-batch ─────────────────────────────────────────────────

class TestProcessBatch:
    def test_valid_edit_returns_docx_and_summary(self, handler, multi_paragraph_docx):
        body = json.dumps(
            {
                "author_name": "Test Author",
                "edits": [
                    {"target_text": "quick brown fox", "new_text": "slow red fox"}
                ],
            }
        )
        resp = handler(
            multipart_request(
                "/api/adeu/process-batch",
                {"file": ("contract.docx", multi_paragraph_docx), "body": body},
            )
        )
        assert resp["statusCode"] == 200
        summary = json.loads(resp["headers"]["X-Batch-Summary"])
        assert summary["applied_edits"] == 1
        assert summary["skipped_edits"] == 0
        assert "contract.docx" in resp["headers"]["Content-Disposition"]
        doc = _docx_body(resp)
        assert len(doc.paragraphs) > 0

    def test_empty_batch_returns_422(self, handler, simple_docx):
        body = json.dumps({"author_name": "Author"})
        resp = handler(
            multipart_request(
                "/api/adeu/process-batch",
                {"file": ("doc.docx", simple_docx), "body": body},
            )
        )
        assert resp["statusCode"] == 422
        assert "At least one edit or action" in _json_body(resp)["detail"]

    def test_unmatched_target_text_rejects_batch(self, handler, simple_docx):
        body = json.dumps(
            {
                "author_name": "Author",
                "edits": [
                    {
                        "target_text": "text that does not exist anywhere",
                        "new_text": "replacement",
                    }
                ],
            }
        )
        resp = handler(
            multipart_request(
                "/api/adeu/process-batch",
                {"file": ("doc.docx", simple_docx), "body": body},
            )
        )
        assert resp["statusCode"] == 422
        assert "detail" in _json_body(resp)


# ── /api/adeu/accept-all ────────────────────────────────────────────────────

class TestAcceptAll:
    def test_accept_all_returns_clean_docx(self, handler, simple_docx):
        resp = handler(
            multipart_request(
                "/api/adeu/accept-all", {"file": ("doc.docx", simple_docx)}
            )
        )
        assert resp["statusCode"] == 200
        doc = _docx_body(resp)
        assert len(doc.paragraphs) > 0

    def test_filename_sanitized_in_content_disposition(self, handler, simple_docx):
        resp = handler(
            multipart_request(
                "/api/adeu/accept-all",
                {"file": ('evil";x=y.docx', simple_docx)},
            )
        )
        assert resp["statusCode"] == 200
        content_disp = resp["headers"]["Content-Disposition"]
        assert '"; x=' not in content_disp
        assert ";" not in content_disp.replace("attachment;", "", 1)


# ── /api/adeu/apply-edits-markdown ──────────────────────────────────────────

class TestApplyEditsMarkdown:
    def test_returns_critic_markup_markdown(self, handler, multi_paragraph_docx):
        body = json.dumps(
            {
                "edits": [
                    {"target_text": "quick brown fox", "new_text": "slow red fox"}
                ]
            }
        )
        resp = handler(
            multipart_request(
                "/api/adeu/apply-edits-markdown",
                {"file": ("doc.docx", multi_paragraph_docx), "body": body},
            )
        )
        assert resp["statusCode"] == 200
        md = _json_body(resp)["markdown"]
        assert "{--" in md or "{++" in md


# ── /api/adeu/diff ──────────────────────────────────────────────────────────

class TestDiff:
    def test_identical_files_have_no_differences(self, handler, simple_docx):
        resp = handler(
            multipart_request(
                "/api/adeu/diff",
                {
                    "original": ("a.docx", simple_docx),
                    "modified": ("b.docx", simple_docx),
                    "compare_clean": "true",
                },
            )
        )
        assert resp["statusCode"] == 200
        body = _json_body(resp)
        assert body["has_differences"] is False

    def test_different_files_have_differences(self, handler):
        def make(text: str) -> bytes:
            doc = Document()
            doc.add_paragraph(text)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        resp = handler(
            multipart_request(
                "/api/adeu/diff",
                {
                    "original": ("a.docx", make("Original text content.")),
                    "modified": ("b.docx", make("Modified text content.")),
                },
            )
        )
        assert resp["statusCode"] == 200
        body = _json_body(resp)
        assert body["has_differences"] is True
        assert len(body["diff"]) > 0


# ── Routing and guards ──────────────────────────────────────────────────────

class TestRoutingAndGuards:
    def test_unknown_route_returns_404(self, handler):
        resp = handler({"path": "/api/adeu/nope", "headers": {}, "body": ""})
        assert resp["statusCode"] == 404

    def test_oversized_content_length_returns_413(self, handler):
        resp = handler(
            {
                "path": "/api/adeu/read",
                "headers": {"content-length": str(51 * 1024 * 1024)},
                "body": "",
            }
        )
        assert resp["statusCode"] == 413
        assert _json_body(resp)["detail"] == "File too large"
