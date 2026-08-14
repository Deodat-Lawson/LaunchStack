"""Shared fixtures for the document-editor service tests.

The suite uses the REAL app from ``app.main`` (routes, trace-id middleware,
and the actual /health handler — the sidecar suite duplicated /health in its
conftest, which was a drift risk). The lifespan is never run (TestClient is
not used as a context manager); the adeu routes are stateless so nothing on
``app.state`` is needed.
"""

import io

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.main import app

# Test API key used for all authenticated requests
TEST_API_KEY = "test-document-editor-api-key"


@pytest.fixture(autouse=True)
def set_api_key_env(monkeypatch):
    """Set DOCUMENT_EDITOR_API_KEY for all tests so the auth dependency passes.
    SIDECAR_API_KEY is cleared so tests exercise the primary variable."""
    monkeypatch.delenv("SIDECAR_API_KEY", raising=False)
    monkeypatch.setenv("DOCUMENT_EDITOR_API_KEY", TEST_API_KEY)


@pytest.fixture
def client():
    """FastAPI test client pre-configured with the API key header."""
    test_client = TestClient(app)
    # Inject the API key header into all requests by default
    test_client.headers.update({"X-API-Key": TEST_API_KEY})
    return test_client


@pytest.fixture
def client_no_auth():
    """Test client with no default API key header attached."""
    return TestClient(app)


@pytest.fixture
def simple_docx() -> bytes:
    """A minimal valid DOCX with known text content."""
    doc = Document()
    doc.add_paragraph("Hello world. This is a test document.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def multi_paragraph_docx() -> bytes:
    """A DOCX with multiple paragraphs for edit testing."""
    doc = Document()
    doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
    doc.add_paragraph("Pack my box with five dozen liquor jugs.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
