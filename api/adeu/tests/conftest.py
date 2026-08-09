"""Fixtures for the deprecated api/adeu Vercel handler tests.

`api/adeu` is not a package (Vercel imports index.py as a standalone file),
so the module is loaded from its file path. Importing it executes the
sys.path bootstrap that pulls the shared schemas from
services/document-editor — the tests verify that wiring explicitly.

Requires the `adeu` pip package (same as services/document-editor).
"""

import base64
import importlib.util
import io
import sys
from pathlib import Path

import pytest
from docx import Document

_INDEX_PATH = Path(__file__).resolve().parent.parent / "index.py"

BOUNDARY = "testboundary1234567890abcdef"


def _load_index_module():
    spec = importlib.util.spec_from_file_location("api_adeu_index", _INDEX_PATH)
    module = importlib.util.module_from_spec(spec)
    sys.modules["api_adeu_index"] = module
    spec.loader.exec_module(module)
    return module


@pytest.fixture(scope="session")
def adeu_index():
    """The loaded api/adeu/index.py module (session-scoped: importing it
    mutates sys.path, do it once)."""
    return _load_index_module()


@pytest.fixture(scope="session")
def handler(adeu_index):
    return adeu_index.handler


def build_multipart_body(fields: dict) -> bytes:
    """Build a multipart/form-data body.

    ``fields`` maps field name -> str (text field) or (filename, bytes)
    tuple (file field).
    """
    chunks: list[bytes] = []
    for name, value in fields.items():
        if isinstance(value, tuple):
            filename, data = value
            chunks.append(
                (
                    f"--{BOUNDARY}\r\n"
                    f'Content-Disposition: form-data; name="{name}"; '
                    f'filename="{filename}"\r\n'
                    "Content-Type: application/octet-stream\r\n\r\n"
                ).encode("utf-8")
                + data
                + b"\r\n"
            )
        else:
            chunks.append(
                (
                    f"--{BOUNDARY}\r\n"
                    f'Content-Disposition: form-data; name="{name}"\r\n\r\n'
                    f"{value}\r\n"
                ).encode("utf-8")
            )
    chunks.append(f"--{BOUNDARY}--\r\n".encode("utf-8"))
    return b"".join(chunks)


def multipart_request(path: str, fields: dict) -> dict:
    """Build a Vercel-style request dict carrying a multipart body."""
    body = build_multipart_body(fields)
    return {
        "path": path,
        "headers": {
            "content-type": f"multipart/form-data; boundary={BOUNDARY}",
            "content-length": str(len(body)),
        },
        "body": base64.b64encode(body).decode("ascii"),
        "isBase64Encoded": True,
    }


@pytest.fixture
def simple_docx() -> bytes:
    """A minimal valid DOCX with known text content."""
    doc = Document()
    doc.add_paragraph("Hello world. This is a test document.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def multi_paragraph_docx() -> bytes:
    """A DOCX with multiple paragraphs for edit testing."""
    doc = Document()
    doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
    doc.add_paragraph("Pack my box with five dozen liquor jugs.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
